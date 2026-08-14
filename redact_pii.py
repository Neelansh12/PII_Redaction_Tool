import re
import io
from typing import Dict, List, Tuple, Generator, Union
import docx
from faker import Faker
import spacy

# Initialize NLP model & Faker
nlp = spacy.load("en_core_web_sm")
fake = Faker("en_IN")
fake.seed_instance(42)  # For deterministic fake data generation

# ---------------------------------------------------------------------------
# Static lookup data
# ---------------------------------------------------------------------------

# Suffixes used by the spaCy-ORG validator (_is_valid_org)
CORPORATE_SUFFIXES = {
    "ltd", "limited", "pvt", "llp", "inc", "co", "company", "corporation",
    "corp", "trust", "gmbh", "sa", "plc", "group", "associates", "partners"
}

# Corporate-suffix PHRASES used by the dedicated regex company detector.
# Kept separate from CORPORATE_SUFFIXES because these are multi-word / punctuated
# phrases that need to be matched literally (in either case) at the end of a
# run of capitalized words.
COMPANY_SUFFIX_ALT = (
    r'(?:Private\s+Limited|Pvt\.?\s*Ltd\.?|Public\s+Limited|Limited|Ltd\.?|'
    r'LLP|L\.L\.P\.?|Inc\.?|Corporation|Corp\.?|PLC)'
)

# Words that should never be treated as the start of a company name even
# though they are capitalized (sentence-initial nouns, generic headings, etc.)
COMPANY_PREFIX_STOPWORDS = {
    "the", "our", "its", "a", "an", "and", "or", "of", "for", "company",
    "offer", "formerly", "investor", "syndicate", "members", "sponsor",
    "banks", "account", "public", "escrow", "collection", "book", "running",
    "lead", "managers", "registrar", "legal", "counsel", "registered",
    "corporate", "office",
}

# Blacklist of terms to never redact (common corporate, regulatory, financial,
# technical, and formatting nouns)
BLACKLIST = {
    # Financial / legal phrases
    "mutual funds", "net proceeds", "qualified institutional buyers", "green shoe option",
    "pat cagr", "cagr", "anchor investor portion", "anchor investors", "offer price",
    "offer structure", "objects of the offer", "promoter group", "promoter selling shareholder",
    "board of directors", "annual report", "draft red herring prospectus", "red herring prospectus",
    "offer letters", "face value", "equity shares",
    # Regulatory & Standards
    "sebi", "bse", "nse", "rbi", "mca", "roc", "cin", "pan", "din", "aadhaar",
    "ssn", "passport", "visa", "companies act", "table", "section", "act", "acts",
    "regulations", "rules", "guidelines", "circulars", "notifications", "gaap",
    "non-gaap", "measures", "measure", "accounting", "standard", "standards",
    "ind as", "eps", "cogs", "inr", "usd", "eur", "euro", "sek", "krona", "kronor",
    "rupee", "rupees", "dollar", "dollars",
    # Common Nouns & Corporate roles
    "board", "company", "companies", "director", "directors", "promoter", "promoters",
    "chairman", "secretary", "auditor", "auditors", "manager", "managers", "officer",
    "officers", "employee", "employees", "employer", "employers", "client", "clients",
    "customer", "customers", "vendor", "vendors", "supplier", "suppliers", "partner",
    "partners", "shareholders", "shareholder", "investors", "investor", "underwriters",
    "underwriter", "lead", "managers", "manager", "counsel", "counsels", "bankers",
    "banker", "allotment", "allot", "allotted", "bids", "bid", "bidder", "bidders",
    "closing", "opening", "date", "day", "days", "annexure", "schedule", "appendix",
    "chapter", "part", "report", "statements", "statement", "definitions", "currency",
    "time", "date", "year", "month", "abbreviations", "prospectus", "draft",
    "red herring prospectus", "herring", "red", "key", "page", "pages", "index",
    "contents", "summary", "introduction", "history", "office", "registered office",
    "corporate office", "registered", "corporate", "facility", "facilities", "plant",
    "plants", "factory", "factories", "unit", "units", "site", "sites", "project",
    "projects", "management", "personnel", "key", "managerial", "executive",
    "independent", "non-executive", "statutory", "legal", "law", "laws",
    "jurisdiction", "court", "tribunal", "arbitrator", "arbitration", "dispute",
    "disputes", "litigation", "proceedings", "claim", "claims", "liability",
    "indemnity", "warranties", "representations", "covenants", "agreements",
    "agreement", "contract", "contracts", "arrangement", "arrangements",
    "transactions", "transaction", "parties", "party", "third", "related",
    "material", "significant", "important", "key", "critical", "major", "minor",
    "principal", "sole", "exclusive", "general", "special", "ordinary", "extraordinary",
    "annual", "interim", "half-yearly", "quarterly", "monthly", "weekly", "daily",
    "hourly", "current", "non-current", "short-term", "long-term", "financial",
    "operational", "commercial", "technical", "administrative", "organizational",
    "competitor", "competitors", "competitiveness", "offered", "shares", "selling",
    "shareholder", "public", "offer", "equity", "preference", "debentures", "bonds",
    "bond", "security", "securities", "capital", "structure", "working", "cash",
    "flow", "balance", "sheet", "assets", "asset", "liabilities", "liability",
    "revenue", "income", "profit", "loss", "expense", "expenses", "tax", "taxes",
    "gst", "tds", "audit", "compliance", "regulatory", "authority", "exchanges",
    "registrar", "registrars", "governing", "governments", "government", "ministry",
    "department", "state", "central", "union", "republic", "national", "international",
    "domestic", "foreign", "local", "primary", "secondary", "offering", "pricing",
    "allotment", "basis", "scheme", "schemes", "portion", "portions", "allocation",
    # Technical & Product Terms
    "asba", "forms", "form", "slip", "slips", "account", "accounts", "association",
    "articles", "aoa", "moa", "conditioning", "air conditioning", "conductors",
    "conductor", "storage", "battery", "energy", "system", "systems", "circuit",
    "kilometers", "kilometer", "transmission", "distribution", "power", "grid",
    "electricity", "wire", "cable", "cables", "copper", "aluminum", "metal",
    "extrusion", "insulation", "continuous", "transposed", "substation", "generator",
    "generators", "turbine", "turbines", "solar", "wind", "renewable", "industry",
    "market", "sector", "business", "quality", "control", "services", "service",
    "product", "products", "development", "research", "reference", "rate", "rates",
    "history", "matters", "change", "changes", "capacity", "expansion", "phase",
    "transit", "marine", "insurance", "cargo", "freight", "logistics", "warehouse",
    "warehousing", "transport", "transportation", "shipping", "delivery", "supply",
    "chain", "demand", "sales", "purchase", "purchases", "acquisition", "merger",
    "amalgamation", "reconstruction", "resolution", "resolutions", "meeting",
    "meetings", "consent", "consents", "approval", "approvals", "order", "orders",
    "certificate", "certificates", "incorporation", "registration", "credential",
    "credentials", "licence", "licences", "license", "licenses", "permit", "permits",
    "clearance", "clearances", "environmental", "pollution", "safety", "health",
    "welfare", "labor", "labour", "employment", "industrial", "disputes", "standing",
    "orders", "provident", "fund", "gratuity", "pension", "bonus", "wages", "salary",
    "compensation", "remuneration", "perquisites", "allowances", "allowance",
    "reimbursement", "reimbursements",
    # Common English words
    "pursuant", "pursuant to", "the", "our", "their", "this", "that", "these",
    "those", "with", "from", "into", "through", "about", "against", "during",
    "before", "after", "above", "below", "under", "over", "between", "among",
    "other", "others", "some", "any", "every", "all", "each", "both", "either",
    "neither", "one", "two", "three", "four", "five", "six", "seven", "eight",
    "nine", "ten", "first", "second", "third", "fourth", "fifth", "last",
    "internal", "external", "risks", "risk", "factors", "factor",
    # Places (to prevent general geographical leakage)
    "india", "maharashtra", "pune", "mumbai", "bombay", "delhi", "bengaluru",
    "chennai", "kolkata", "hyderabad", "ahmedabad", "gurugram", "noida",
    "khed", "chakan", "taluka", "village", "birdewadi", "ahilyanagar", "ahmednagar",
    "gujarat", "uae", "sweden", "us", "u.s.", "united states", "europe", "european"
}

# Honorifics to validate single-word PERSON names
HONORIFICS = {"mr", "mrs", "ms", "dr", "shri", "smt"}

UPPER_NAME_RE = re.compile(r'\b(?:[A-Z]{2,}(?:\s+[A-Z]{2,}){1,5})\b')
PERSON_NAME_RE = re.compile(r'\b(?:[A-Z][a-zA-Z\'.-]+(?:\s+(?:[A-Z][a-zA-Z\'.-]+|[A-Z]{2,})){1,5})\b')

# Role-based person name detector
ROLE_NAME_RE = re.compile(
    r'\b(?i:Contact\s+Person|Chairman|Managing\s+Director|Whole[- ]?time\s+Director|'
    r'Executive\s+Director|Independent\s+Director|Director|CEO|Chief\s+Executive\s+Officer|'
    r'CFO|Chief\s+Financial\s+Officer|Company\s+Secretary|Promoter|Promoters)'
    r'(?:\s+(?:and|&)\s+(?i:compliance\s+officer|executive\s+director))?'
    r'(?:\s+(?i:of\s+(?:our|the|its)\s+Company))?'
    r'\s*(?:,\s*)?(?:[:\-]|(?i:is\s+being|is\s+our|is|are|being\s+our|being|namely|of|our|being\s+the))?\s*'
    r'([A-Z][a-zA-Z\.\']+(?:\s+[A-Z][a-zA-Z\.\']+){1,3})\b'
)

# Honorific-based person name detector
HONORIFIC_NAME_RE = re.compile(
    r'\b(?:Mr\.?|Mrs\.?|Dr\.?|Shri|Smt|Ms\.?)\s+'
    r'([A-Z][a-zA-Z\.\']+(?:\s+[A-Z][a-zA-Z\.\']+){1,3})\b'
)

# Labels that introduce a physical / mailing address.
ADDRESS_LABEL_RE = re.compile(
    r'\b(?:Registered\s+and\s+Corporate\s+Office|Registered\s+Office|Corporate\s+Office|'
    r'Registered\s+Address|Corporate\s+Address|Correspondence\s+Address|'
    r'Mailing\s+Address|Factory|Plant|Works|Located\s+at|Residing\s+at|'
    r'Situated\s+at|Office|Address)'
    r'\s*(?:of\s+(?:our|the|its)\s+Company\s*)?(?:is\s+situated\s+at\s*|at\s*|:\s*)',
    re.IGNORECASE
)

# Words that mark the end of an address run within a paragraph
ADDRESS_STOP_WORDS = (
    r'(?:Telephone|Tel\.?|Phone|Fax|E-?mail|Website|Contact\s?Person|CIN\b|PAN\b|'
    r'Registration\s+Number|SEBI|Investor\s+grievance)'
)
ADDRESS_STOP_RE = re.compile(r'\b' + ADDRESS_STOP_WORDS + r'\b', re.IGNORECASE)
ADDRESS_TRAILING_CONNECTOR_RE = re.compile(r'\s*(?:and\s+its|and\s+the|and)\s*$', re.IGNORECASE)

# Structural (label-free) address patterns
ADDRESS_ANCHOR_RE = re.compile(
    r'\b(?:Gat|Survey|Sy|Plot|Flat|House|H\.?\s?No)\s*\.?\s*No\.?\s*[:.]?\s*'
    r'[\dA-Za-z][\dA-Za-z/,&\.\-\s\u2013\u2014()]{0,120}?'
    r'(?=(?:\s*\.\s+[A-Z])|(?:\s*\b' + ADDRESS_STOP_WORDS + r'\b)|$)',
    re.IGNORECASE
)
ADDRESS_CITY_PIN_RE = re.compile(
    r'\b[A-Z][a-zA-Z]+(?:\s*\([A-Za-z]+\))?[\s]{0,2}[\u2013\-]?[\s]{0,2}\d{3}\s?\d{3}\b'
    r'(?:,?\s*(?!' + ADDRESS_STOP_WORDS + r')(?:Maharashtra|India|[A-Z][a-zA-Z]+))*'
)
ADDRESS_KEYWORDS_RE = re.compile(
    r'\b(?:village|vill\.?|taluka|tehsil|district|dist\.?|gat\s*no|survey\s*no|'
    r'sy\.?\s*no|plot\s*no|road|street|marg|nagar|chowk|sector|phase|'
    r'industrial\s+area|colony|lane|wing|floor|tower|building|bldg|block|flat)\b',
    re.IGNORECASE
)
ADDRESS_NUM_LEAD_RE = re.compile(
    r'^\s*\d[\dA-Za-z/,&\.\-\s\u2013\u2014()]{2,100}?'
    r'(?=(?:\s*\.\s+[A-Z])|(?:\s*\b' + ADDRESS_STOP_WORDS + r'\b)|$)'
)

ADDRESS_BLOCK_RE = re.compile(
    r'(?i)(?:^|[,;\n])\s*(?:\d{1,5}\s*(?:[-–]\s*\d{1,5})?|(?:Flat|House|Plot|Survey|Bldg|Building|Wing|Floor|Shop|Unit|Office|Road|Street|Lane|Marg|Sector|Phase|Village|Taluka|District|Colony|Nagar)[A-Za-z0-9 .,&/-]*|[A-Z][A-Za-z0-9 .,&/-]*)'
    r'(?:[^\n;]{0,220}?)(?:\b\d{6}\b|\b(?:Maharashtra|Karnataka|Tamil\s+Nadu|Gujarat|Delhi|Mumbai|Pune|Bengaluru|Hyderabad|Ahmedabad|India)\b)'
    r'(?:[^\n;]{0,180})'
)

# Date of birth context detector
DOB_DETECTOR_RE = re.compile(
    r'\b(?:date\s+of\s+birth|d\.?o\.?b\.?|birth\s*date|born\s+on)\s*[:\-]?\s*'
    r'('
    r'\d{1,2}(?:st|nd|rd|th)?[-\/\.]\d{1,2}(?:st|nd|rd|th)?[-\/\.]\d{2,4}|'
    r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4}|'
    r'\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}'
    r')',
    re.IGNORECASE
)

# Labels that introduce a phone number
PHONE_LABEL_RE = re.compile(r'(?:telephone|tel|phone|mobile|mob|fax|ph)\s*\.?\s*(?:no\.?|number)?\s*:?\s*$', re.IGNORECASE)

# Improved phone regex
PHONE_CANDIDATE_RE = re.compile(
    r'(?<!\d)(?:\+\s?\d{1,4}[-\s]?)?\(?\d{2,6}\)?[-\s]?\d{2,6}(?:[-\s]?\d{2,6}){1,4}(?!\d)'
)

# Priority for PII types when resolving overlaps of EQUAL-LENGTH spans.
TYPE_PRIORITY = {
    "CREDIT_CARD": 85, "SSN": 80, "EMAIL": 75, "PHONE": 70, "IP_ADDRESS": 65,
    "DATE_OF_BIRTH": 60, "ADDRESS": 55, "ORG": 50, "PERSON": 45,
}


class PIIRedactor:
    def __init__(self):
        # Cache replacement mapping. Keyed by (entity_type, original_text).
        self.mapping: Dict[Tuple[str, str], str] = {}
        self._canonical_mapping: Dict[Tuple[str, str], str] = {}

        self.regex_patterns = {
            "EMAIL": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',
            "SSN": r'\b\d{3}-\d{2}-\d{4}\b',
            "CREDIT_CARD": r'\b(?:\d[ -]*?){13,19}\b',
            "IP_ADDRESS": r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
        }
        self.company_regex = re.compile(
            r'\b(?:[A-Z0-9][A-Za-z0-9&\.\'-]*\s+(?:(?:and|of|for|in|&)\s+)?){1,8}(?i:' + COMPANY_SUFFIX_ALT + r')\b'
        )

    # ------------------------------------------------------------------
    # Validation helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _normalize_for_key(value: str) -> str:
        value = value.strip('.,;:!?"\'()[]{}')
        value = re.sub(r'\s+', ' ', value.strip())
        return value.casefold()

    def _mapping_key(self, entity_type: str, original_text: str) -> Tuple[str, str]:
        return entity_type, self._normalize_for_key(original_text)

    def is_luhn_valid(self, number: str) -> bool:
        """Validates a card number using the Luhn algorithm."""
        digits = [int(c) for c in number if c.isdigit()]
        if not 13 <= len(digits) <= 19:
            return False

        total = 0
        reversed_digits = digits[::-1]
        for i, d in enumerate(reversed_digits):
            if i % 2 == 1:
                d *= 2
                if d > 9:
                    d -= 9
            total += d
        return total % 10 == 0

    def _is_valid_person(self, text: str, context: str = "") -> bool:
        """Validates if a PERSON entity is a real personal name, not noise."""
        cleaned = text.strip().strip('.,;:!?"\'()[]{}')
        lower_cleaned = cleaned.lower()

        if lower_cleaned in BLACKLIST:
            return False

        cleaned_words = [w for w in re.split(r'\W+', cleaned) if w]
        if not cleaned_words:
            return False

        if any(w.lower() in BLACKLIST for w in cleaned_words):
            return False

        org_terms = {"infra", "park", "private", "limited", "llp", "industrial", "motors", "logistics", "group", "trust",
                     "capital", "services", "international", "parked", "company", "corporation", "partners", "associates",
                     "industries", "holdings", "foundation"}
        if any(w.lower() in org_terms for w in cleaned_words):
            return False

        if any(ch.isdigit() for ch in cleaned):
            return False

        if len(cleaned_words) >= 2:
            valid_tokens = [w for w in cleaned_words if w.lower() not in {"and", "of", "&"}]
            if len(valid_tokens) >= 2:
                capitalized_words = [w for w in valid_tokens if w[0].isupper() or w.isupper()]
                if len(capitalized_words) >= 2:
                    return True
                if all(w.isupper() for w in valid_tokens):
                    return True

        if context:
            pos = context.find(text)
            if pos != -1:
                before = context[:pos].strip().lower()
                after = context[pos + len(text):].strip().lower()
                for title in HONORIFICS:
                    if before.endswith(title) or before.endswith(title + "."):
                        return True
                    if after.startswith(title) or after.startswith(title + "."):
                        return True

        return False

    def _trim_and_validate_person(self, text: str, start_char: int, end_char: int, context: str = "") -> Tuple[int, int, str, bool]:
        """Trims leading honorifics and trailing role/corporate designations from a PERSON entity before validating."""
        cleaned = text.strip()
        words = cleaned.split()
        if not words:
            return start_char, end_char, "", False

        # Trim leading honorifics
        honorifics_to_trim = {"mr", "mrs", "ms", "dr", "shri", "smt"}
        while words and words[0].lower().strip('.,;:-') in honorifics_to_trim:
            words.pop(0)

        # Trim trailing role/corporate words
        trim_words = {
            "company", "secretary", "director", "directors", "promoter", "promoters",
            "chairman", "cfo", "ceo", "kmp", "kmps", "sm", "sms", "executive", "officer", "compliance"
        }
        while words and words[-1].lower().strip('.,;:-') in trim_words:
            words.pop()

        if not words:
            return start_char, end_char, "", False

        new_text = " ".join(words).rstrip('.,;:- ')
        pos = text.find(new_text)
        if pos != -1:
            new_start = start_char + pos
            new_end = new_start + len(new_text)
        else:
            new_start = start_char
            new_end = start_char + len(new_text)

        is_valid = self._is_valid_person(new_text, context)
        return new_start, new_end, new_text, is_valid

    def _is_valid_org(self, text: str) -> bool:
        """Validates if a spaCy ORG entity is a real organization, not general text."""
        cleaned = text.strip().strip('.,;:!?"\'()[]{}')
        lower_cleaned = cleaned.lower()

        if lower_cleaned in BLACKLIST:
            return False

        cleaned_words = [w for w in re.split(r'\W+', cleaned) if w]
        if not cleaned_words:
            return False

        if any(w.lower() in BLACKLIST for w in cleaned_words):
            return False

        if self._is_valid_person(cleaned):
            legal_terms = {'private', 'limited', 'ltd', 'llp', 'inc', 'corp', 'corporation', 'trust', 'group', 'company', 'partners', 'associates', 'industries', 'services', 'infrastructure', 'holdings'}
            has_legal_suffix = any(w.lower() in legal_terms for w in cleaned_words)
            if not has_legal_suffix:
                return False

        if any(w.lower() in {"green", "shoe", "option", "mutual", "funds", "net", "proceeds",
                             "qualified", "institutional", "buyers", "cagr", "infra", "park"}
               for w in cleaned_words):
            return False

        if len(cleaned_words) == 1 and lower_cleaned in BLACKLIST:
            return False

        has_suffix = any(w.lower() in CORPORATE_SUFFIXES for w in cleaned_words)
        has_num = any(ch.isdigit() for ch in cleaned)
        if has_suffix:
            return True
        if has_num and len(cleaned_words) >= 2:
            return True

        if len(cleaned_words) >= 2:
            capitalized_words = [w for w in cleaned_words if w[0].isupper() or w.isupper()]
            if len(capitalized_words) >= 2:
                return True

        return False

    def _is_address_continuation(self, text: str) -> bool:
        """Determines if a paragraph line is a continuation of an address block (e.g. 'Maharashtra, India')."""
        cleaned = text.strip().strip(';,.- ')
        if not cleaned:
            return False
        words = [w.lower().strip(';,.-') for w in cleaned.split() if w]
        if not words:
            return False
        if len(cleaned) > 60:
            return False

        address_continuation_words = {
            "india", "maharashtra", "pune", "mumbai", "bombay", "delhi", "bengaluru",
            "chennai", "kolkata", "hyderabad", "ahmedabad", "gurugram", "noida",
            "khed", "chakan", "taluka", "village", "district", "dist", "state", "country",
            "gujarat", "karnataka", "goa", "telangana", "andhra", "pradesh", "tamil", "nadu"
        }

        for w in words:
            if w in {"and", "or", "of", "in", "to", "near", "opposite", "behind", "next"}:
                continue
            if w in address_continuation_words:
                continue
            if w in BLACKLIST:
                continue
            return False

        return True

    def _find_company_spans(self, text: str) -> List[Tuple[int, int, str, str, bool]]:
        """Dedicated regex layer for organization / company names."""
        spans = []
        for m in self.company_regex.finditer(text):
            start, end = m.start(), m.end()
            candidate = m.group(0).strip()
            words = candidate.split()

            trim = 0
            while trim < len(words) - 1 and words[trim].lower().strip('.,') in COMPANY_PREFIX_STOPWORDS:
                trim += 1
            words = words[trim:]

            candidate = ' '.join(words)
            if not candidate or candidate.lower() in BLACKLIST:
                continue

            if any(tok.lower() in BLACKLIST for tok in re.split(r'\W+', candidate) if tok):
                continue

            if any(ch.isdigit() for ch in candidate) and 'private' not in candidate.lower() and 'limited' not in candidate.lower():
                # Allow numeric parts such as 'KSH Infra Park 5 Private Limited' while still rejecting generic headings
                pass

            non_suffix_words = [w for w in words if not re.fullmatch(COMPANY_SUFFIX_ALT, w, re.IGNORECASE)]
            if len(non_suffix_words) < 2:
                continue

            spans.append((start, end, "ORG", candidate, True))
        return spans

    def _find_address_spans(self, text: str) -> List[Tuple[int, int, str, str, bool]]:
        """Detects physical / mailing addresses using label-anchored and structural strategies."""
        spans: List[Tuple[int, int, str, str, bool]] = []

        # Strategy 1: explicit label
        for m in ADDRESS_LABEL_RE.finditer(text):
            start = m.end()
            next_label = ADDRESS_LABEL_RE.search(text, start)
            limit = next_label.start() if next_label else len(text)
            stop_m = ADDRESS_STOP_RE.search(text, start, limit)
            end = stop_m.start() if stop_m else limit

            segment = text[start:end]
            sent_end = re.search(r'\.\s+[A-Z]', segment)
            if sent_end:
                end = start + sent_end.start() + 1

            raw = text[start:end]
            raw = ADDRESS_TRAILING_CONNECTOR_RE.sub('', raw)
            candidate = raw.strip().rstrip(';,. ')
            if candidate and any(c.isdigit() for c in candidate) and len(candidate) > 8:
                actual_end = start + len(candidate)
                spans.append((start, actual_end, "ADDRESS", candidate, True))

        # Strategy 2: structural
        for m in ADDRESS_ANCHOR_RE.finditer(text):
            candidate = m.group(0).strip()
            if candidate:
                spans.append((m.start(), m.start() + len(candidate), "ADDRESS", candidate, True))

        for m in ADDRESS_CITY_PIN_RE.finditer(text):
            spans.append((m.start(), m.end(), "ADDRESS", m.group(0), True))

        for m in ADDRESS_BLOCK_RE.finditer(text):
            candidate = m.group(0).strip().strip(';, ')
            if candidate and any(ch.isdigit() for ch in candidate) and (re.search(r'\d{6}', candidate) or ADDRESS_KEYWORDS_RE.search(candidate)):
                spans.append((m.start(), m.end(), "ADDRESS", candidate, True))

        m = ADDRESS_NUM_LEAD_RE.match(text)
        if m and ADDRESS_KEYWORDS_RE.search(m.group(0)) and any(c.isalpha() for c in m.group(0)):
            candidate = m.group(0).strip()
            spans.append((m.start(), m.start() + len(candidate), "ADDRESS", candidate, True))

        return spans

    def _detect_phone_spans(self, text: str) -> List[Tuple[int, int, str, str, bool]]:
        """Detects telephone numbers in Indian and international formats."""
        spans = []
        for m in PHONE_CANDIDATE_RE.finditer(text):
            raw = m.group(0)
            digits = re.sub(r'\D', '', raw)
            if not (8 <= len(digits) <= 13):
                continue

            stripped = raw.strip()
            has_plus = stripped.startswith('+')
            has_sep = bool(re.search(r'[-\s()]', stripped))

            # Reject fiscal-year false positives
            groups = [g for g in re.split(r'[-\s()]+', stripped) if g]
            if len(groups) == 2 and all(len(g) == 4 for g in groups):
                try:
                    y1, y2 = int(groups[0]), int(groups[1])
                    if 1900 <= y1 <= 2099 and 1900 <= y2 <= 2099 and 0 <= y2 - y1 <= 1:
                        continue
                except ValueError:
                    pass

            if not has_plus:
                if not has_sep:
                    before = text[max(0, m.start() - 25):m.start()]
                    if not PHONE_LABEL_RE.search(before):
                        continue

            spans.append((m.start(), m.end(), "PHONE", raw, True))
        return spans

    # ------------------------------------------------------------------
    # Fake replacement generation
    # ------------------------------------------------------------------

    def _get_fake_replacement(self, entity_type: str, original_text: str) -> str:
        """Returns or creates a synthetic alternative for detected PII."""
        canonical_key = self._mapping_key(entity_type, original_text)
        if canonical_key in self._canonical_mapping:
            return self._canonical_mapping[canonical_key]

        replacement = original_text
        if entity_type == "EMAIL":
            replacement = fake.email()
        elif entity_type == "PHONE":
            digits = ''.join(str(fake.random_int(min=0, max=9)) for _ in range(10))
            replacement = f"+91 {digits[:2]} {digits[2:6]} {digits[6:]}"
        elif entity_type == "SSN":
            replacement = fake.ssn()
        elif entity_type == "CREDIT_CARD":
            replacement = fake.credit_card_number()
        elif entity_type == "IP_ADDRESS":
            replacement = fake.ipv4()
        elif entity_type == "DATE_OF_BIRTH":
            replacement = fake.date_of_birth().strftime("%d/%m/%Y")
        elif entity_type == "PERSON":
            replacement = fake.name()
        elif entity_type == "ORG":
            replacement = fake.company()
        elif entity_type == "ADDRESS":
            replacement = fake.address().replace("\n", ", ")

        self.mapping[(entity_type, original_text)] = replacement
        self._canonical_mapping[canonical_key] = replacement
        return replacement

    # ------------------------------------------------------------------
    # DOCX traversal
    # ------------------------------------------------------------------

    def _iter_paragraphs(self, doc: docx.Document) -> Generator[docx.text.paragraph.Paragraph, None, None]:
        """Recursively yields all paragraphs from the body, tables, and headers/footers."""
        for p in doc.paragraphs:
            yield p

        for table in doc.tables:
            yield from self._iter_table_paragraphs(table)

        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    yield p
                for table in section.header.tables:
                    yield from self._iter_table_paragraphs(table)
            if section.footer:
                for p in section.footer.paragraphs:
                    yield p
                for table in section.footer.tables:
                    yield from self._iter_table_paragraphs(table)

    def _iter_table_paragraphs(self, table) -> Generator[docx.text.paragraph.Paragraph, None, None]:
        """Recursively yields paragraphs from within a table's cells, handling nested tables."""
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested_table in cell.tables:
                    yield from self._iter_table_paragraphs(nested_table)

    # ------------------------------------------------------------------
    # Detection pipeline
    # ------------------------------------------------------------------

    def _detect_pii_spans(self, text: str) -> List[Tuple[int, int, str, str, bool]]:
        """Detects all PII occurrences in the text using regex, dedicated structural detectors, and spaCy NER."""
        spans = []

        # 1. Generic structured regex spans
        for pii_type, pattern in self.regex_patterns.items():
            for match in re.finditer(pattern, text):
                match_str = match.group(0)
                start, end = match.start(), match.end()

                # Prevent matching sub-segments of longer digit sequences
                if pii_type in ["CREDIT_CARD", "SSN"]:
                    left = start - 1
                    is_sub = False
                    while left >= 0:
                        char = text[left]
                        if char.isdigit():
                            is_sub = True
                            break
                        elif char in [' ', '-']:
                            left -= 1
                        else:
                            break

                    right = end
                    while right < len(text) and not is_sub:
                        char = text[right]
                        if char.isdigit():
                            is_sub = True
                            break
                        elif char in [' ', '-']:
                            right += 1
                        else:
                            break

                    if is_sub:
                        continue

                # Apply Luhn check filter for credit cards
                if pii_type == "CREDIT_CARD" and not self.is_luhn_valid(match_str):
                    continue

                spans.append((start, end, pii_type, match_str, True))

        # 2. Dedicated contextual DOB detector
        for match in DOB_DETECTOR_RE.finditer(text):
            start, end = match.start(1), match.end(1)
            match_str = match.group(1)
            spans.append((start, end, "DATE_OF_BIRTH", match_str, True))

        # 3. Dedicated PHONE detector
        spans.extend(self._detect_phone_spans(text))

        # 4. Dedicated company-name regex layer
        spans.extend(self._find_company_spans(text))

        # 5. Dedicated address detectors
        spans.extend(self._find_address_spans(text))

        # 6. Person names (Role-anchored, title-anchored, and all-caps names)
        for match in ROLE_NAME_RE.finditer(text):
            name = match.group(1).strip()
            if self._is_valid_person(name, context=text):
                spans.append((match.start(1), match.end(1), "PERSON", name, True))

        for match in HONORIFIC_NAME_RE.finditer(text):
            name = match.group(1).strip()
            if self._is_valid_person(name, context=text):
                spans.append((match.start(1), match.end(1), "PERSON", name, True))

        for match in UPPER_NAME_RE.finditer(text):
            name = match.group(0).strip()
            if len(name.split()) >= 2 and self._is_valid_person(name, context=text):
                spans.append((match.start(), match.end(), "PERSON", name, True))

        for match in PERSON_NAME_RE.finditer(text):
            name = match.group(0).strip()
            if len(name.split()) >= 2 and self._is_valid_person(name, context=text):
                spans.append((match.start(), match.end(), "PERSON", name, True))

        # 7. SpaCy NER spans (filtered to prevent false positives)
        doc = nlp(text)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                start_p, end_p, name_p, is_v = self._trim_and_validate_person(ent.text, ent.start_char, ent.end_char, context=text)
                if is_v:
                    spans.append((start_p, end_p, "PERSON", name_p, False))
            elif ent.label_ in ["ORG", "COMPANY"]:
                if self._is_valid_org(ent.text):
                    spans.append((ent.start_char, ent.end_char, "ORG", ent.text, False))
            elif ent.label_ in ["GPE", "LOC", "FAC"]:
                cleaned = ent.text.strip()
                if any(c.isdigit() for c in cleaned) and cleaned.lower() not in BLACKLIST:
                    spans.append((ent.start_char, ent.end_char, "ADDRESS", cleaned, False))

        return spans

    def _resolve_conflicts(self, spans: List[Tuple[int, int, str, str, bool]]) -> List[Tuple[int, int, str, str]]:
        """Resolves overlapping PII spans based on structured priority and length."""
        def sort_key(span):
            start, end, pii_type, _, is_structured = span
            length = end - start
            return (1 if is_structured else 0, length, TYPE_PRIORITY.get(pii_type, 0), start)

        sorted_spans = sorted(spans, key=sort_key, reverse=True)

        selected_spans = []
        covered_indices = set()

        for start, end, pii_type, original_text, is_structured in sorted_spans:
            if any(i in covered_indices for i in range(start, end)):
                continue
            selected_spans.append((start, end, pii_type, original_text))
            for i in range(start, end):
                covered_indices.add(i)

        return sorted(selected_spans, key=lambda x: x[0])

    # ------------------------------------------------------------------
    # Replacement application
    # ------------------------------------------------------------------

    def _apply_spans_to_paragraph(self, paragraph, spans: List[Tuple[int, int, str, str]]) -> None:
        """Replaces resolved PII spans directly within the paragraph's runs so that formatting is preserved."""
        if not spans:
            return

        runs = paragraph.runs
        run_texts = [r.text for r in runs]
        if not runs or ''.join(run_texts) != paragraph.text:
            # Fallback: simple string-level replacement on the whole paragraph
            text = paragraph.text
            for start, end, pii_type, original_text in sorted(spans, key=lambda s: s[0], reverse=True):
                replacement = self._canonical_mapping.get(self._mapping_key(pii_type, original_text))
                if replacement is not None:
                    text = text[:start] + replacement + text[end:]
            paragraph.text = text
            return

        boundaries = []
        pos = 0
        for t in run_texts:
            boundaries.append((pos, pos + len(t)))
            pos += len(t)

        for start, end, pii_type, original_text in sorted(spans, key=lambda s: s[0], reverse=True):
            replacement = self._canonical_mapping.get(self._mapping_key(pii_type, original_text))
            if replacement is None:
                continue

            overlapping = [i for i, (rs, re_) in enumerate(boundaries) if re_ > start and rs < end]
            if not overlapping:
                continue

            first_idx, last_idx = overlapping[0], overlapping[-1]
            rs0, _ = boundaries[first_idx]
            local_start = start - rs0

            if first_idx == last_idx:
                t = runs[first_idx].text
                local_end = end - rs0
                runs[first_idx].text = t[:local_start] + replacement + t[local_end:]
            else:
                t0 = runs[first_idx].text
                runs[first_idx].text = t0[:local_start] + replacement
                for mid in overlapping[1:-1]:
                    runs[mid].text = ''
                rsl, _ = boundaries[last_idx]
                local_end_last = end - rsl
                tl = runs[last_idx].text
                runs[last_idx].text = tl[local_end_last:]

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def detect_and_replace_text(self, text: str) -> str:
        """Runs full detection + replacement on a single plain-text string."""
        spans = self._detect_pii_spans(text)
        resolved_spans = self._resolve_conflicts(spans)

        for start, end, pii_type, original_text in resolved_spans:
            if self._mapping_key(pii_type, original_text) not in self._canonical_mapping:
                self._get_fake_replacement(pii_type, original_text)

        new_text = text
        for start, end, pii_type, original_text in sorted(resolved_spans, key=lambda s: s[0], reverse=True):
            replacement = self._canonical_mapping.get(self._mapping_key(pii_type, original_text))
            if replacement is not None:
                new_text = new_text[:start] + replacement + new_text[end:]
        return new_text

    def redact_document(self, input_file: Union[str, io.BytesIO], output_file: Union[str, io.BytesIO]) -> Tuple[Dict[str, int], Dict[Tuple[str, str], str]]:
        """Performs a two-stage redaction on the given DOCX file (path or file-like object)."""
        doc = docx.Document(input_file)

        self.mapping = {}
        self._canonical_mapping = {}
        counts: Dict[str, int] = {}

        paragraphs = list(self._iter_paragraphs(doc))

        # --- PASS 1: Detect PII and populate self.mapping + detect address block continuations ---
        prev_was_address = False
        
        for p in paragraphs:
            text = p.text
            if not text.strip():
                prev_was_address = False
                continue

            spans = self._detect_pii_spans(text)
            
            # Check for context-aware address block continuation
            has_address = any(s[2] == "ADDRESS" for s in spans)
            if not has_address and prev_was_address and self._is_address_continuation(text):
                # Classify the entire paragraph text as ADDRESS
                spans.append((0, len(text), "ADDRESS", text, True))
                has_address = True

            resolved_spans = self._resolve_conflicts(spans)
            for start, end, pii_type, original_text in resolved_spans:
                if self._mapping_key(pii_type, original_text) not in self._canonical_mapping:
                    self._get_fake_replacement(pii_type, original_text)
             
            prev_was_address = has_address

        # --- PASS 2: Apply replacements to each paragraph (including literal searches for all known PII) ---
        prev_was_address = False
        
        for p in paragraphs:
            text = p.text
            if not text.strip():
                prev_was_address = False
                continue

            spans = self._detect_pii_spans(text)
            
            # Check for context-aware address block continuation
            has_address = any(s[2] == "ADDRESS" for s in spans)
            if not has_address and prev_was_address and self._is_address_continuation(text):
                spans.append((0, len(text), "ADDRESS", text, True))
                has_address = True

            # Also search for all globally mapped PII texts to guarantee coverage
            for (pii_type, original_text) in list(self.mapping.keys()):
                start = 0
                while True:
                    pos = text.find(original_text, start)
                    if pos == -1:
                        break
                    spans.append((pos, pos + len(original_text), pii_type, original_text, True))
                    start = pos + 1

            resolved_spans = self._resolve_conflicts(spans)
            
            # Count the final redacted occurrences
            for start, end, pii_type, original_text in resolved_spans:
                counts[pii_type] = counts.get(pii_type, 0) + 1
            
            self._apply_spans_to_paragraph(p, resolved_spans)
            prev_was_address = has_address

        doc.save(output_file)
        return counts, self.mapping

    def redact_docx(self, input_filepath: str, output_filepath: str):
        """Reads DOCX, redacts text in paragraphs and tables, and saves output."""
        self.redact_document(input_filepath, output_filepath)
        print(f"Redacted document saved successfully to: {output_filepath}")

    # ------------------------------------------------------------------
    # Evaluation
    # ------------------------------------------------------------------

    def evaluate(self) -> Dict[str, Dict[str, object]]:
        """Small precision/recall-style evaluation covering all 9 required PII categories plus negative examples."""
        positive_cases = {
            "PERSON": "Contact Person: Sarthak Malvadkar, Company Secretary.",
            "EMAIL": "E-mail: cs.connect@kshinternational.com; Website: www.kshinternational.com",
            "PHONE": "Telephone: + 91 20 45053237",
            "ORG": "Our Company was originally incorporated as Bhandary Metal Extrusion Private Limited.",
            "ADDRESS": "Registered Office: 11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed, Pune \u2013 410 501, Maharashtra, India;",
            "SSN": "SSN on file: 123-45-6789.",
            "CREDIT_CARD": "Card number 4532015112830366 was used for the transaction.",
            "DATE_OF_BIRTH": "Date of Birth: 12/05/1998",
            "IP_ADDRESS": "The server responded from 192.168.1.100.",
        }
        negative_cases = {
            "ordinary_date": "Dated December 10, 2025, the Board approved the resolution.",
            "bare_city": "Our facilities are located in Pune, Maharashtra, India.",
            "fiscal_year_range": "Revenue grew between Fiscals 2022-2023 and 2023-2024.",
            "generic_corporate_noun": "The Board of Directors approved the Annual Report.",
        }

        results: Dict[str, Dict[str, object]] = {"positive": {}, "negative": {}}
        for label, text in positive_cases.items():
            redacted = self.detect_and_replace_text(text)
            results["positive"][label] = {
                "input": text,
                "output": redacted,
                "changed": redacted != text,
            }
        for label, text in negative_cases.items():
            redacted = self.detect_and_replace_text(text)
            results["negative"][label] = {
                "input": text,
                "output": redacted,
                "unchanged": redacted == text,
            }
        return results
